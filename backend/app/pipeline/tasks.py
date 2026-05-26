import json
from datetime import datetime, timedelta, timezone

import redis
import sentry_sdk
from sqlalchemy import select

from app.core.config import settings
from app.core.logging import get_logger
from app.db.session import SyncSessionLocal
from app.models.scenario import Scenario
from app.models.session import SimulationSession, SessionDecision
from app.pipeline.celery_app import celery_app
from app.pipeline.claude_client import extract_scenario_from_document, generate_debrief_report
from app.pipeline.ingestion import (
    search_cisa_advisories,
    fetch_plain_text,
    search_sec_8k_filings,
    fetch_sec_filing_text,
    fetch_hhs_breach_csv,
    fetch_rss_article_urls,
    is_source_already_processed,
    RSS_FEEDS,
    HHS_BREACH_CSV_URL,
    fetch_cisa_kev_text,
)
from app.pipeline.embeddings import generate_embedding, scenario_text


logger = get_logger(__name__)
redis_client = redis.Redis.from_url(settings.REDIS_URL, decode_responses=True)


def _store_task_failure(task_name: str, task_id: str, error: str) -> None:
    payload = {
        "task_name": task_name,
        "task_id": task_id,
        "error": error,
        "timestamp": datetime.now(timezone.utc).isoformat(),
    }
    redis_client.setex(f"task_failure:{task_name}:{task_id}", 86400, json.dumps(payload))


def _handle_task_failure(task, error: Exception) -> None:
    task_name = task.name
    task_id = task.request.id
    logger.exception(
        "Celery task failed",
        extra={
            "task_name": task_name,
            "task_id": task_id,
            "task_args": task.request.args,
            "task_kwargs": task.request.kwargs,
            "exception_type": type(error).__name__,
        },
    )
    if settings.SENTRY_DSN:
        sentry_sdk.capture_exception(error)
    try:
        _store_task_failure(task_name, task_id, str(error))
    except Exception:
        logger.exception("Failed to store Celery task failure", extra={"task_name": task_name, "task_id": task_id})


@celery_app.task(bind=True, name="app.pipeline.tasks.ingest_cisa_advisories")
def ingest_cisa_advisories(self, limit: int = 10, days_back: int = 90):
    try:
        advisory_urls = search_cisa_advisories(limit=limit, days_back=days_back)
        results = []
        for url in advisory_urls:
            try:
                result = process_advisory_url.delay(url, source_type="cisa", source_reference=url.split("/")[-1])
                results.append({"url": url, "task_id": result.id})
            except Exception as e:
                results.append({"url": url, "error": str(e)})
        return results
    except Exception as e:
        _handle_task_failure(self, e)
        raise


@celery_app.task(bind=True, name="app.pipeline.tasks.process_advisory_url")
def process_advisory_url(self, url: str, source_type: str = "cisa", source_reference: str = None, preloaded_text: str = None):
    try:
        if source_reference and is_source_already_processed(source_reference):
            return {"status": "skipped", "reason": "Already processed", "source_reference": source_reference}

        # Handle CISA KEV pseudo-URLs (cisa-kev://CVE-XXXX-XXXX)
        if url.startswith("cisa-kev://"):
            cve_id = url[len("cisa-kev://"):]
            text = preloaded_text or fetch_cisa_kev_text(cve_id)
        else:
            text = preloaded_text or fetch_plain_text(url)
        if not text or len(text) < 200:
            raise ValueError(f"Insufficient content from {url}")

        extracted = extract_scenario_from_document(text)

        if extracted.get("extraction_confidence", 0) < 0.4:
            return {"status": "rejected", "reason": "Low extraction confidence", "confidence": extracted.get("extraction_confidence")}

        scenario_data = {
            "title": extracted.get("title", "Untitled Scenario"),
            "source_type": source_type,
            "source_url": url,
            "source_reference": source_reference,
            "incident_date": extracted.get("incident_date"),
            "incident_duration_hours": extracted.get("incident_duration_hours"),
            "initial_access_vector": extracted.get("initial_access_vector"),
            "industry_vertical": extracted.get("industry_vertical"),
            "affected_asset_types": extracted.get("affected_asset_types"),
            "mitre_techniques": extracted.get("mitre_techniques"),
            "nist_controls": extracted.get("nist_controls"),
            "regulatory_frameworks": extracted.get("regulatory_frameworks"),
            "extraction_confidence": extracted.get("extraction_confidence"),
            "alert_sequence": extracted.get("alert_sequence"),
            "decision_tree": extracted.get("decision_tree"),
            "status": "review" if extracted.get("extraction_confidence", 0) >= 0.7 else "draft",
        }

        _save_scenario_sync(scenario_data)
        return {"status": "success", "title": scenario_data["title"], "confidence": extracted.get("extraction_confidence")}
    except Exception as e:
        _handle_task_failure(self, e)
        raise


@celery_app.task(bind=True, name="app.pipeline.tasks.ingest_sec_8k_filings")
def ingest_sec_8k_filings(self, days_back: int = 30, limit: int = 5):
    """Search SEC EDGAR for recent cybersecurity-related 8-K filings and queue extraction."""
    try:
        end_date = datetime.now(timezone.utc).date()
        start_date = end_date - timedelta(days=days_back)
        filings = search_sec_8k_filings(
            start_date=start_date.isoformat(),
            end_date=end_date.isoformat(),
            limit=limit,
        )
        results = []
        for filing in filings:
            source_ref = filing["accession_no"]
            if not source_ref:
                continue
            try:
                text = fetch_sec_filing_text(filing["entity_id"], filing["accession_no"])
                if not text:
                    results.append({"filing": source_ref, "status": "skipped", "reason": "No text extracted"})
                    continue
                result = process_advisory_url.delay(
                    url=f"https://www.sec.gov/cgi-bin/browse-edgar?action=getcompany&CIK={filing['entity_id']}&type=8-K",
                    source_type="sec_8k",
                    source_reference=source_ref,
                    preloaded_text=text,
                )
                results.append({"filing": source_ref, "entity": filing["entity_name"], "task_id": result.id})
            except Exception as e:
                results.append({"filing": source_ref, "error": str(e)})
        return results
    except Exception as e:
        _handle_task_failure(self, e)
        raise


@celery_app.task(bind=True, name="app.pipeline.tasks.ingest_hhs_breaches")
def ingest_hhs_breaches(self, min_individuals: int = 10000, limit: int = 5):
    """Download HHS OCR Breach Portal CSV and queue significant healthcare breaches for extraction."""
    try:
        breaches = fetch_hhs_breach_csv(min_individuals=min_individuals)
        results = []
        for breach in breaches[:limit]:
            source_ref = breach["source_reference"]
            try:
                result = process_advisory_url.delay(
                    url=HHS_BREACH_CSV_URL,
                    source_type="hhs",
                    source_reference=source_ref,
                    preloaded_text=breach["summary"],
                )
                results.append({"entity": breach["name"], "individuals": breach["individuals"], "task_id": result.id})
            except Exception as e:
                results.append({"entity": breach["name"], "error": str(e)})
        return results
    except Exception as e:
        _handle_task_failure(self, e)
        raise


@celery_app.task(bind=True, name="app.pipeline.tasks.ingest_rss_feeds")
def ingest_rss_feeds(self, limit_per_feed: int = 3):
    """Fetch Krebs on Security and SANS ISC RSS feeds and queue articles for extraction."""
    try:
        results = []
        for feed_name, feed_url in RSS_FEEDS.items():
            urls = fetch_rss_article_urls(feed_url, limit=limit_per_feed)
            for url in urls:
                source_ref = f"{feed_name}-{url.rstrip('/').split('/')[-1][:80]}"
                try:
                    result = process_advisory_url.delay(
                        url=url,
                        source_type="cisa",
                        source_reference=source_ref,
                    )
                    results.append({"feed": feed_name, "url": url, "task_id": result.id})
                except Exception as e:
                    results.append({"feed": feed_name, "url": url, "error": str(e)})
        return results
    except Exception as e:
        _handle_task_failure(self, e)
        raise


@celery_app.task(bind=True, name="app.pipeline.tasks.generate_session_debrief")
def generate_session_debrief(self, session_id: str):
    try:
        _generate_debrief_sync(session_id)
    except Exception as e:
        _handle_task_failure(self, e)
        raise


def _save_scenario_sync(data: dict) -> None:
    with SyncSessionLocal() as db:
        if data.get("source_reference"):
            existing = db.execute(
                select(Scenario).where(Scenario.source_reference == data["source_reference"])
            ).scalar_one_or_none()
            if existing:
                logger.info("Scenario already exists, skipping", extra={"source_reference": data["source_reference"]})
                return
        scenario = Scenario(**{k: v for k, v in data.items() if v is not None})
        try:
            scenario.embedding = generate_embedding(
                scenario_text(
                    title=data.get("title", ""),
                    initial_access_vector=data.get("initial_access_vector"),
                    mitre_techniques=data.get("mitre_techniques"),
                    nist_controls=data.get("nist_controls"),
                    industry_vertical=data.get("industry_vertical"),
                    regulatory_frameworks=data.get("regulatory_frameworks"),
                )
            )
        except Exception:
            logger.exception("Embedding generation failed — saving scenario without embedding")
        db.add(scenario)
        db.commit()


@celery_app.task(bind=True, name="app.pipeline.tasks.backfill_scenario_embeddings")
def backfill_scenario_embeddings(self):
    """Generate embeddings for existing scenarios that have none. Run once after migration."""
    with SyncSessionLocal() as db:
        scenarios = db.execute(
            select(Scenario).where(Scenario.embedding == None)  # noqa: E711
        ).scalars().all()
        updated = 0
        for s in scenarios:
            try:
                s.embedding = generate_embedding(
                    scenario_text(
                        title=s.title,
                        initial_access_vector=s.initial_access_vector,
                        mitre_techniques=s.mitre_techniques,
                        nist_controls=s.nist_controls,
                        industry_vertical=s.industry_vertical,
                        regulatory_frameworks=s.regulatory_frameworks,
                    )
                )
                updated += 1
            except Exception:
                logger.exception("Embedding generation failed for scenario", extra={"scenario_id": s.id})
        db.commit()
        return {"backfilled": updated, "total": len(scenarios)}


def _generate_debrief_sync(session_id: str) -> None:
    with SyncSessionLocal() as db:
        session = db.execute(
            select(SimulationSession).where(SimulationSession.id == session_id)
        ).scalar_one_or_none()
        if not session:
            return

        scenario = db.execute(
            select(Scenario).where(Scenario.id == session.scenario_id)
        ).scalar_one_or_none()

        decisions_raw = db.execute(
            select(SessionDecision).where(SessionDecision.session_id == session_id)
        ).scalars().all()

        decision_tree_map = {g["id"]: g for g in (scenario.decision_tree or [])}

        decisions = []
        control_gaps = []
        for d in decisions_raw:
            gate = decision_tree_map.get(d.decision_gate_id, {})
            options = gate.get("options", [])
            decisions.append({
                "gate_id": d.decision_gate_id,
                "team_choice": options[d.chosen_option_index]["text"] if d.chosen_option_index < len(options) else "Unknown",
                "correct_choice": options[gate.get("correct_index", 0)]["text"] if options else "Unknown",
                "is_correct": d.is_correct,
                "nist_ref": d.nist_control_ref,
                "mitre_technique": d.mitre_technique,
            })
            if not d.is_correct and d.nist_control_ref:
                control_gaps.append({"control": d.nist_control_ref, "gate_id": d.decision_gate_id})

        report = generate_debrief_report(
            scenario_title=scenario.title,
            source_reference=scenario.source_reference,
            score=session.team_score or 0,
            correct=session.decisions_correct,
            total=session.decisions_made,
            decisions=decisions,
            control_gaps=control_gaps,
        )

        session.debrief_report = report
        session.debrief_generated_at = datetime.utcnow()
        db.commit()


@celery_app.task(bind=True, name="app.pipeline.tasks.send_weekly_slack_snippet")
def send_weekly_slack_snippet(self):
    """Pick a random approved public scenario and post a snippet to the Slack channel."""
    from app.services.slack_service import send_webhook_message, build_scenario_snippet_blocks

    try:
        with SyncSessionLocal() as db:
            scenario = db.execute(
                select(Scenario)
                .where(Scenario.status == "approved", Scenario.is_private == False)  # noqa: E712
                .order_by(Scenario.play_count.asc())
                .limit(1)
            ).scalar_one_or_none()

        if not scenario:
            logger.info("No approved public scenarios available for weekly Slack snippet")
            return {"status": "skipped", "reason": "no scenarios"}

        blocks = build_scenario_snippet_blocks({
            "title": scenario.title,
            "difficulty": scenario.difficulty,
            "industry_vertical": scenario.industry_vertical,
            "mitre_techniques": scenario.mitre_techniques,
            "nist_controls": scenario.nist_controls,
            "estimated_minutes": scenario.estimated_minutes,
        })
        send_webhook_message(text=f"Weekly Scenario: {scenario.title}", blocks=blocks)
        return {"status": "sent", "scenario_id": scenario.id}
    except Exception as e:
        _handle_task_failure(self, e)
        raise


@celery_app.task(bind=True, name="app.pipeline.tasks.process_uploaded_document_task")
def process_uploaded_document_task(self, document_id: str):
    from pypdf import PdfReader
    from app.models.breach_document import BreachDocument

    with SyncSessionLocal() as db:
        doc = db.execute(
            select(BreachDocument).where(BreachDocument.id == document_id)
        ).scalar_one_or_none()
        if not doc:
            logger.error("BreachDocument %s not found inside database", document_id)
            return

        try:
            # 1. Parse content
            file_key = doc.file_key
            if file_key.startswith("s3://"):
                import io
                import boto3
                from app.core.config import settings

                # Parse bucket and key from s3://bucket/key
                s3_parts = file_key[5:].split("/", 1)
                bucket_name = s3_parts[0]
                object_key = s3_parts[1]

                s3_client = boto3.client(
                    "s3",
                    aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
                    aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY,
                    region_name=settings.AWS_REGION
                )
                response = s3_client.get_object(Bucket=bucket_name, Key=object_key)
                content_bytes = response["Body"].read()

                if file_key.lower().endswith(".pdf"):
                    stream = io.BytesIO(content_bytes)
                    reader = PdfReader(stream)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() or ""
                elif file_key.lower().endswith(".docx"):
                    import io as _io
                    from docx import Document
                    doc_stream = _io.BytesIO(content_bytes)
                    docx_doc = Document(doc_stream)
                    text = "\n".join(p.text for p in docx_doc.paragraphs)
                else:
                    text = content_bytes.decode("utf-8", errors="ignore")
            else:
                if file_key.lower().endswith(".pdf"):
                    reader = PdfReader(file_key)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() or ""
                elif file_key.lower().endswith(".docx"):
                    from docx import Document
                    docx_doc = Document(file_key)
                    text = "\n".join(p.text for p in docx_doc.paragraphs)
                else:
                    with open(file_key, "r", encoding="utf-8", errors="ignore") as f:
                        text = f.read()

            if not text or len(text) < 200:
                raise ValueError("Parsed document has insufficient text content (< 200 characters)")

            # 2. Extract Scenario from document using Claude
            extracted = extract_scenario_from_document(text)

            # 3. Create Scenario
            scenario_data = {
                "title": extracted.get("title", doc.filename),
                "source_type": "private",
                "source_reference": doc.filename,
                "incident_date": extracted.get("incident_date"),
                "incident_duration_hours": extracted.get("incident_duration_hours"),
                "initial_access_vector": extracted.get("initial_access_vector"),
                "industry_vertical": extracted.get("industry_vertical"),
                "affected_asset_types": extracted.get("affected_asset_types"),
                "mitre_techniques": extracted.get("mitre_techniques"),
                "nist_controls": extracted.get("nist_controls"),
                "regulatory_frameworks": extracted.get("regulatory_frameworks"),
                "extraction_confidence": extracted.get("extraction_confidence"),
                "alert_sequence": extracted.get("alert_sequence"),
                "decision_tree": extracted.get("decision_tree"),
                "status": "review" if extracted.get("extraction_confidence", 0) >= 0.7 else "draft",
                "is_private": True,
                "owner_org_id": doc.organization_id,
            }

            scenario = Scenario(**{k: v for k, v in scenario_data.items() if v is not None})
            db.add(scenario)
            db.flush()

            # Update BreachDocument
            doc.extracted_scenario_id = scenario.id
            doc.status = "completed"
            db.commit()

            return {"status": "success", "scenario_id": scenario.id}

        except Exception as e:
            db.rollback()
            doc.status = "failed"
            db.commit()
            _handle_task_failure(self, e)
            raise
